"""
AegisLang L1 Ingestion Layer

Purpose: Intake, clean, and normalize policy text from heterogeneous sources.

Functional Requirements:
- ING-001: Parse PDF documents with text extraction
- ING-002: Parse DOCX documents preserving structure
- ING-003: Parse Markdown files
- ING-004: Parse HTML from web-based policy portals
- ING-005: Apply OCR for scanned documents (optional)
- ING-006: Chunk text into semantically coherent sections
- ING-007: Detect and preserve document hierarchy (sections, subsections)
- ING-008: Emit standardized JSON for downstream agents
"""

from __future__ import annotations

import hashlib
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import structlog
import tiktoken
from pydantic import BaseModel, Field

logger = structlog.get_logger(__name__)


# -----------------------------------------------------------------------------
# Schema Definitions
# -----------------------------------------------------------------------------


class TextChunk(BaseModel):
    """A single chunk of text within a section."""

    chunk_id: str = Field(..., description="Unique chunk identifier")
    text: str = Field(..., description="Chunk text content")
    token_count: int = Field(..., description="Number of tokens in chunk")
    embedding_vector: list[float] | None = Field(
        default=None, description="Optional pre-computed embedding"
    )


class DocumentSection(BaseModel):
    """A section within the document hierarchy."""

    section_id: str = Field(..., description="Unique section identifier")
    section_title: str = Field(..., description="Section title or heading")
    parent_section: str | None = Field(
        default=None, description="Parent section ID if nested"
    )
    hierarchy_level: int = Field(
        ..., ge=1, le=6, description="Heading level (1-6)"
    )
    text_chunks: list[TextChunk] = Field(
        default_factory=list, description="Text chunks in this section"
    )


class DocumentMetadata(BaseModel):
    """Metadata about the ingested document."""

    source_file: str = Field(..., description="Original file path or URL")
    ingestion_timestamp: str = Field(
        ..., description="ISO 8601 timestamp of ingestion"
    )
    document_type: str = Field(
        ..., description="Document format: pdf, docx, markdown, html"
    )
    page_count: int | None = Field(default=None, description="Number of pages")
    language: str = Field(default="en", description="Detected language code")
    hash: str = Field(..., description="SHA-256 hash of source document")


class IngestedDocument(BaseModel):
    """Output schema for the ingestion layer."""

    doc_id: str = Field(
        ..., pattern=r"^[A-Z0-9_]+$", description="Unique document identifier"
    )
    metadata: DocumentMetadata = Field(..., description="Document metadata")
    sections: list[DocumentSection] = Field(
        default_factory=list, description="Document sections with chunks"
    )


# -----------------------------------------------------------------------------
# Chunking Configuration
# -----------------------------------------------------------------------------


class ChunkingConfig(BaseModel):
    """Configuration for text chunking."""

    target_tokens: int = Field(default=768, description="Target chunk size")
    min_tokens: int = Field(default=256, description="Minimum chunk size")
    max_tokens: int = Field(default=1024, description="Maximum chunk size")
    overlap_tokens: int = Field(
        default=64, description="Overlap between chunks"
    )


# -----------------------------------------------------------------------------
# Text Chunker
# -----------------------------------------------------------------------------


class SemanticChunker:
    """Chunks text into semantically coherent sections."""

    def __init__(self, config: ChunkingConfig | None = None):
        self.config = config or ChunkingConfig()
        self._tokenizer = None
        try:
            self._tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception as e:
            logger.warning("tiktoken_fallback", reason=str(e))
            # Will use character-based estimation

    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken or fallback estimation."""
        if self._tokenizer:
            return len(self._tokenizer.encode(text))
        # Fallback: estimate ~4 chars per token (average for English)
        return len(text) // 4

    def chunk_text(
        self, text: str, section_id: str
    ) -> list[TextChunk]:
        """
        Split text into chunks based on semantic boundaries.

        Uses paragraph breaks as primary split points, then sentence
        boundaries for finer control.
        """
        if not text.strip():
            return []

        # Split by paragraph first
        paragraphs = re.split(r"\n\s*\n", text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks: list[TextChunk] = []
        current_chunk: list[str] = []
        current_tokens = 0

        for para in paragraphs:
            para_tokens = self.count_tokens(para)

            # If single paragraph exceeds max, split by sentences
            if para_tokens > self.config.max_tokens:
                # Flush current chunk
                if current_chunk:
                    chunks.append(
                        self._create_chunk(current_chunk, section_id, len(chunks))
                    )
                    current_chunk = []
                    current_tokens = 0

                # Split large paragraph by sentences
                sentence_chunks = self._split_by_sentences(para, section_id, len(chunks))
                chunks.extend(sentence_chunks)
                continue

            # Check if adding this paragraph exceeds target
            if current_tokens + para_tokens > self.config.target_tokens:
                # Flush if we have minimum tokens
                if current_tokens >= self.config.min_tokens:
                    chunks.append(
                        self._create_chunk(current_chunk, section_id, len(chunks))
                    )
                    # Add overlap from previous chunk
                    overlap_text = self._get_overlap_text(current_chunk)
                    current_chunk = [overlap_text] if overlap_text else []
                    current_tokens = self.count_tokens(overlap_text) if overlap_text else 0

            current_chunk.append(para)
            current_tokens += para_tokens

        # Flush remaining
        if current_chunk:
            chunks.append(
                self._create_chunk(current_chunk, section_id, len(chunks))
            )

        return chunks

    def _split_by_sentences(
        self, text: str, section_id: str, chunk_offset: int
    ) -> list[TextChunk]:
        """Split text by sentence boundaries."""
        # Simple sentence splitting regex
        sentences = re.split(r"(?<=[.!?])\s+", text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks: list[TextChunk] = []
        current_sentences: list[str] = []
        current_tokens = 0

        for sentence in sentences:
            sent_tokens = self.count_tokens(sentence)

            if current_tokens + sent_tokens > self.config.target_tokens:
                if current_sentences:
                    chunks.append(
                        self._create_chunk(
                            current_sentences,
                            section_id,
                            chunk_offset + len(chunks),
                        )
                    )
                    current_sentences = []
                    current_tokens = 0

            current_sentences.append(sentence)
            current_tokens += sent_tokens

        if current_sentences:
            chunks.append(
                self._create_chunk(
                    current_sentences, section_id, chunk_offset + len(chunks)
                )
            )

        return chunks

    def _create_chunk(
        self, parts: list[str], section_id: str, index: int
    ) -> TextChunk:
        """Create a TextChunk from text parts."""
        text = " ".join(parts)
        return TextChunk(
            chunk_id=f"{section_id}_C{index:03d}",
            text=text,
            token_count=self.count_tokens(text),
        )

    def _get_overlap_text(self, parts: list[str]) -> str:
        """Get overlap text from the end of parts."""
        if not parts:
            return ""

        # Take last paragraph or portion thereof
        last_text = parts[-1]
        tokens = self._tokenizer.encode(last_text)

        if len(tokens) <= self.config.overlap_tokens:
            return last_text

        # Take last N tokens
        overlap_tokens = tokens[-self.config.overlap_tokens :]
        return self._tokenizer.decode(overlap_tokens)


# -----------------------------------------------------------------------------
# Document Parsers
# -----------------------------------------------------------------------------


class BaseDocumentParser:
    """Base class for document parsers."""

    def __init__(self):
        self.chunker = SemanticChunker()

    def parse(self, file_path: Path) -> tuple[list[DocumentSection], int | None]:
        """
        Parse document and return sections with optional page count.

        Returns:
            Tuple of (sections, page_count)
        """
        raise NotImplementedError

    def compute_hash(self, file_path: Path) -> str:
        """Compute SHA-256 hash of file."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _extract_hierarchy(
        self, text: str, doc_id: str
    ) -> list[DocumentSection]:
        """
        Extract document hierarchy from markdown-style headings.

        Detects # headings and creates section structure.
        """
        sections: list[DocumentSection] = []
        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

        # Find all headings with their positions
        headings = [
            (m.start(), len(m.group(1)), m.group(2).strip())
            for m in heading_pattern.finditer(text)
        ]

        if not headings:
            # No headings found, treat entire text as single section
            section = DocumentSection(
                section_id=f"{doc_id}_S001",
                section_title="Document Content",
                hierarchy_level=1,
                text_chunks=self.chunker.chunk_text(text, f"{doc_id}_S001"),
            )
            return [section] if section.text_chunks else []

        # Build section hierarchy
        parent_stack: list[tuple[int, str]] = []  # (level, section_id)

        for i, (pos, level, title) in enumerate(headings):
            # Get text until next heading or end
            if i + 1 < len(headings):
                end_pos = headings[i + 1][0]
            else:
                end_pos = len(text)

            # Extract section text (skip the heading line itself)
            heading_end = text.find("\n", pos)
            if heading_end == -1:
                heading_end = pos
            section_text = text[heading_end:end_pos].strip()

            # Generate section ID
            section_id = f"{doc_id}_S{i + 1:03d}"

            # Find parent
            while parent_stack and parent_stack[-1][0] >= level:
                parent_stack.pop()

            parent_id = parent_stack[-1][1] if parent_stack else None

            # Create section
            section = DocumentSection(
                section_id=section_id,
                section_title=title,
                parent_section=parent_id,
                hierarchy_level=level,
                text_chunks=self.chunker.chunk_text(section_text, section_id),
            )

            sections.append(section)
            parent_stack.append((level, section_id))

        return sections


class PDFParser(BaseDocumentParser):
    """Parser for PDF documents."""

    def parse(self, file_path: Path) -> tuple[list[DocumentSection], int | None]:
        """Parse PDF document."""
        try:
            from pdfminer.high_level import extract_text
            from pdfminer.pdfpage import PDFPage
        except ImportError as e:
            logger.error("pdfminer.six not installed", error=str(e))
            raise ImportError(
                "pdfminer.six is required for PDF parsing. "
                "Install with: pip install pdfminer.six"
            ) from e

        # Extract text
        text = extract_text(str(file_path))

        # Count pages
        page_count = 0
        with open(file_path, "rb") as f:
            for _ in PDFPage.get_pages(f):
                page_count += 1

        # Generate doc_id from filename
        doc_id = self._generate_doc_id(file_path)

        sections = self._extract_hierarchy(text, doc_id)

        return sections, page_count

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate document ID from filename."""
        name = file_path.stem.upper()
        # Replace non-alphanumeric with underscore
        name = re.sub(r"[^A-Z0-9]", "_", name)
        # Remove consecutive underscores
        name = re.sub(r"_+", "_", name)
        return name.strip("_")


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX documents."""

    def parse(self, file_path: Path) -> tuple[list[DocumentSection], int | None]:
        """Parse DOCX document preserving structure."""
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError as e:
            logger.error("python-docx not installed", error=str(e))
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            ) from e

        doc = Document(str(file_path))
        doc_id = self._generate_doc_id(file_path)

        sections: list[DocumentSection] = []
        current_text: list[str] = []
        current_heading: str | None = None
        current_level: int = 1
        section_count = 0
        parent_stack: list[tuple[int, str]] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            # Check if this is a heading
            if style_name.startswith("Heading"):
                # Flush previous section
                if current_text or current_heading:
                    section_count += 1
                    section_id = f"{doc_id}_S{section_count:03d}"

                    # Find parent
                    while parent_stack and parent_stack[-1][0] >= current_level:
                        parent_stack.pop()
                    parent_id = parent_stack[-1][1] if parent_stack else None

                    section = DocumentSection(
                        section_id=section_id,
                        section_title=current_heading or "Untitled Section",
                        parent_section=parent_id,
                        hierarchy_level=current_level,
                        text_chunks=self.chunker.chunk_text(
                            "\n\n".join(current_text), section_id
                        ),
                    )
                    sections.append(section)
                    parent_stack.append((current_level, section_id))

                # Start new section
                try:
                    current_level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    current_level = 1

                current_heading = para.text.strip()
                current_text = []
            else:
                # Regular paragraph
                text = para.text.strip()
                if text:
                    current_text.append(text)

        # Flush final section
        if current_text or current_heading:
            section_count += 1
            section_id = f"{doc_id}_S{section_count:03d}"

            while parent_stack and parent_stack[-1][0] >= current_level:
                parent_stack.pop()
            parent_id = parent_stack[-1][1] if parent_stack else None

            section = DocumentSection(
                section_id=section_id,
                section_title=current_heading or "Document Content",
                parent_section=parent_id,
                hierarchy_level=current_level,
                text_chunks=self.chunker.chunk_text(
                    "\n\n".join(current_text), section_id
                ),
            )
            sections.append(section)

        return sections, None  # DOCX doesn't have traditional pages

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate document ID from filename."""
        name = file_path.stem.upper()
        name = re.sub(r"[^A-Z0-9]", "_", name)
        name = re.sub(r"_+", "_", name)
        return name.strip("_")


class MarkdownParser(BaseDocumentParser):
    """Parser for Markdown documents."""

    def parse(self, file_path: Path) -> tuple[list[DocumentSection], int | None]:
        """Parse Markdown document."""
        text = file_path.read_text(encoding="utf-8")
        doc_id = self._generate_doc_id(file_path)

        sections = self._extract_hierarchy(text, doc_id)

        return sections, None

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate document ID from filename."""
        name = file_path.stem.upper()
        name = re.sub(r"[^A-Z0-9]", "_", name)
        name = re.sub(r"_+", "_", name)
        return name.strip("_")


class HTMLParser(BaseDocumentParser):
    """Parser for HTML documents."""

    def parse(self, file_path: Path) -> tuple[list[DocumentSection], int | None]:
        """Parse HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError as e:
            logger.error("beautifulsoup4 not installed", error=str(e))
            raise ImportError(
                "beautifulsoup4 is required for HTML parsing. "
                "Install with: pip install beautifulsoup4"
            ) from e

        html_content = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        doc_id = self._generate_doc_id(file_path)
        sections: list[DocumentSection] = []
        section_count = 0
        parent_stack: list[tuple[int, str]] = []

        # Find all heading elements
        headings = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])

        if not headings:
            # No headings, extract all text as single section
            text = soup.get_text(separator="\n\n", strip=True)
            section_id = f"{doc_id}_S001"
            section = DocumentSection(
                section_id=section_id,
                section_title="Document Content",
                hierarchy_level=1,
                text_chunks=self.chunker.chunk_text(text, section_id),
            )
            return [section] if section.text_chunks else [], None

        for i, heading in enumerate(headings):
            level = int(heading.name[1])  # h1 -> 1, h2 -> 2, etc.
            title = heading.get_text(strip=True)

            # Get content until next heading
            content_parts: list[str] = []
            for sibling in heading.next_siblings:
                if sibling.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    break
                if hasattr(sibling, "get_text"):
                    text = sibling.get_text(strip=True)
                    if text:
                        content_parts.append(text)

            section_count += 1
            section_id = f"{doc_id}_S{section_count:03d}"

            # Find parent
            while parent_stack and parent_stack[-1][0] >= level:
                parent_stack.pop()
            parent_id = parent_stack[-1][1] if parent_stack else None

            section = DocumentSection(
                section_id=section_id,
                section_title=title,
                parent_section=parent_id,
                hierarchy_level=level,
                text_chunks=self.chunker.chunk_text(
                    "\n\n".join(content_parts), section_id
                ),
            )

            sections.append(section)
            parent_stack.append((level, section_id))

        return sections, None

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate document ID from filename."""
        name = file_path.stem.upper()
        name = re.sub(r"[^A-Z0-9]", "_", name)
        name = re.sub(r"_+", "_", name)
        return name.strip("_")


# -----------------------------------------------------------------------------
# Main Ingestor Agent
# -----------------------------------------------------------------------------


class AegisIngestor:
    """
    L1 Ingestion Layer Agent.

    Intakes policy documents from various formats and produces
    normalized, chunked output for downstream processing.
    """

    SUPPORTED_FORMATS = {
        ".pdf": ("pdf", PDFParser),
        ".docx": ("docx", DOCXParser),
        ".md": ("markdown", MarkdownParser),
        ".markdown": ("markdown", MarkdownParser),
        ".html": ("html", HTMLParser),
        ".htm": ("html", HTMLParser),
    }

    def __init__(self, chunking_config: ChunkingConfig | None = None):
        """Initialize the ingestor with optional chunking configuration."""
        self.chunking_config = chunking_config or ChunkingConfig()
        self._parsers: dict[str, BaseDocumentParser] = {}

    def _get_parser(self, file_ext: str) -> BaseDocumentParser:
        """Get or create parser for file extension."""
        if file_ext not in self._parsers:
            if file_ext not in self.SUPPORTED_FORMATS:
                raise ValueError(f"Unsupported file format: {file_ext}")

            _, parser_class = self.SUPPORTED_FORMATS[file_ext]
            parser = parser_class()
            parser.chunker = SemanticChunker(self.chunking_config)
            self._parsers[file_ext] = parser

        return self._parsers[file_ext]

    def ingest(self, file_path: str | Path) -> IngestedDocument:
        """
        Ingest a document and produce normalized output.

        Args:
            file_path: Path to the document file

        Returns:
            IngestedDocument with sections and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        file_ext = path.suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: {file_ext}. "
                f"Supported: {list(self.SUPPORTED_FORMATS.keys())}"
            )

        logger.info(
            "ingesting_document",
            file_path=str(path),
            format=file_ext,
        )

        # Get parser and process
        parser = self._get_parser(file_ext)
        sections, page_count = parser.parse(path)

        # Compute document hash
        doc_hash = parser.compute_hash(path)

        # Generate document ID
        doc_id = self._generate_doc_id(path)

        # Get document type
        doc_type, _ = self.SUPPORTED_FORMATS[file_ext]

        # Create metadata
        metadata = DocumentMetadata(
            source_file=str(path.absolute()),
            ingestion_timestamp=datetime.now(timezone.utc).isoformat(),
            document_type=doc_type,
            page_count=page_count,
            language="en",  # TODO: Add language detection
            hash=doc_hash,
        )

        # Create output document
        document = IngestedDocument(
            doc_id=doc_id,
            metadata=metadata,
            sections=sections,
        )

        logger.info(
            "ingestion_complete",
            doc_id=doc_id,
            sections=len(sections),
            total_chunks=sum(len(s.text_chunks) for s in sections),
        )

        return document

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID."""
        name = file_path.stem.upper()
        name = re.sub(r"[^A-Z0-9]", "_", name)
        name = re.sub(r"_+", "_", name)
        name = name.strip("_")

        # Add short UUID suffix for uniqueness
        suffix = uuid.uuid4().hex[:6].upper()
        return f"{name}_{suffix}"

    def ingest_to_json(self, file_path: str | Path) -> str:
        """
        Ingest document and return JSON string.

        Convenience method for pipeline integration.
        """
        document = self.ingest(file_path)
        return document.model_dump_json(indent=2)

    def ingest_to_dict(self, file_path: str | Path) -> dict[str, Any]:
        """
        Ingest document and return dictionary.

        Convenience method for pipeline integration.
        """
        document = self.ingest(file_path)
        return document.model_dump()


# -----------------------------------------------------------------------------
# Event Publishing (Agent-OS Integration)
# -----------------------------------------------------------------------------


async def publish_ingested_event(
    document: IngestedDocument,
    redis_url: str | None = None,
) -> None:
    """
    Publish policy.ingested event to Agent-OS event bus.

    Args:
        document: The ingested document
        redis_url: Redis connection URL
    """
    if redis_url is None:
        import os
        redis_url = os.environ.get("REDIS_URL", "redis://localhost:6379/0")

    try:
        import redis.asyncio as redis_async

        client = redis_async.from_url(redis_url)
        await client.publish(
            "policy.ingested",
            document.model_dump_json(),
        )
        await client.aclose()

        logger.info(
            "event_published",
            topic="policy.ingested",
            doc_id=document.doc_id,
        )
    except Exception as e:
        logger.warning(
            "event_publish_failed",
            topic="policy.ingested",
            error=str(e),
        )


# -----------------------------------------------------------------------------
# CLI Entry Point
# -----------------------------------------------------------------------------


def main() -> None:
    """Command-line interface for document ingestion."""
    import argparse
    import json
    import sys

    parser = argparse.ArgumentParser(
        description="AegisLang Document Ingestor - L1 Ingestion Layer"
    )
    parser.add_argument(
        "file",
        help="Path to document file (PDF, DOCX, Markdown, or HTML)",
    )
    parser.add_argument(
        "-o", "--output",
        help="Output file path (default: stdout)",
    )
    parser.add_argument(
        "--target-tokens",
        type=int,
        default=768,
        help="Target chunk size in tokens (default: 768)",
    )
    parser.add_argument(
        "--min-tokens",
        type=int,
        default=256,
        help="Minimum chunk size in tokens (default: 256)",
    )
    parser.add_argument(
        "--max-tokens",
        type=int,
        default=1024,
        help="Maximum chunk size in tokens (default: 1024)",
    )

    args = parser.parse_args()

    # Configure chunking
    config = ChunkingConfig(
        target_tokens=args.target_tokens,
        min_tokens=args.min_tokens,
        max_tokens=args.max_tokens,
    )

    # Run ingestion
    ingestor = AegisIngestor(chunking_config=config)

    try:
        result = ingestor.ingest_to_json(args.file)

        if args.output:
            Path(args.output).write_text(result)
            print(f"Output written to: {args.output}", file=sys.stderr)
        else:
            print(result)

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
